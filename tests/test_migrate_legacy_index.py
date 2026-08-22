import json
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import Mock, patch

from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings

from app.embedding_client import EmbeddingConfig
from app.index_manifest import load_index_manifest
from app.langchain_store import close_vector_store, open_vector_store
from app.material_ingestion import MaterialManager
from app.migrate_legacy_index import (
    CANDIDATE_VECTOR_STORE_DIRNAME,
    ROLLED_BACK_CANDIDATE_DIRNAME,
    LegacyMigrationError,
    MigrationLayout,
    add_staged_to_candidate,
    build_base_candidate,
    plan_migration,
    prepare_migration,
    prepare_staged_snapshot,
    promote_candidate,
    rollback_promotion,
    main,
    validate_candidate,
)


MIGRATION_ID = "a" * 32


class KeywordEmbeddings(Embeddings):
    @staticmethod
    def _embed(text: str) -> list[float]:
        return [1.0, 0.0] if "base" in text.casefold() else [0.0, 1.0]

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [self._embed(text) for text in texts]

    def embed_query(self, text: str) -> list[float]:
        return self._embed(text)


class FailingEmbeddings(KeywordEmbeddings):
    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        raise RuntimeError("simulated embedding failure")


class LegacyIndexMigrationTests(unittest.TestCase):
    def make_layout(self, root: Path) -> MigrationLayout:
        return MigrationLayout(
            documents_dir=root / "documents",
            vector_store_dir=root / "vector_store",
            pending_uploads_dir=root / "pending_uploads",
            pending_deletions_dir=root / "pending_deletions",
            migrations_dir=root / "migrations",
        )

    def prepare_source_state(self, root: Path) -> tuple[MigrationLayout, str]:
        layout = self.make_layout(root)
        layout.documents_dir.mkdir()
        (layout.documents_dir / "base.md").write_text(
            "base material",
            encoding="utf-8",
        )
        legacy = open_vector_store(KeywordEmbeddings(), layout.vector_store_dir)
        legacy.add_documents(
            [Document(page_content="legacy base", metadata={"source": "base.md"})],
            ids=["legacy-id"],
        )
        close_vector_store(legacy)

        manager = MaterialManager(
            documents_dir=layout.documents_dir,
            pending_uploads_dir=layout.pending_uploads_dir,
            pending_deletions_dir=layout.pending_deletions_dir,
            sync_index=lambda chunks: None,
            estimate_index_batches=lambda chunks: 1,
        )
        staged = manager.stage_upload(
            filename="word-notes.docx",
            content_type="application/zip",
            stream=BytesIO(self.docx_bytes()),
        )
        return layout, staged.upload_id

    @staticmethod
    def docx_bytes() -> bytes:
        from docx import Document as WordDocument

        document = WordDocument()
        document.add_paragraph("staged word material")
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def config() -> EmbeddingConfig:
        return EmbeddingConfig(
            api_key="test-key",
            base_url="https://example.com/v1",
            model="text-embedding-v4",
            dimensions=1024,
        )

    @staticmethod
    def record_ids(path: Path) -> set[str]:
        store = open_vector_store(None, path)
        try:
            return set(store.get(include=[])["ids"])
        finally:
            close_vector_store(store)

    def build_complete_candidate(
        self,
        layout: MigrationLayout,
        upload_id: str,
    ) -> Path:
        workspace = prepare_migration(
            layout,
            [upload_id],
            migration_id=MIGRATION_ID,
        )
        build_base_candidate(
            workspace,
            confirm_api_cost=True,
            confirm_batches=1,
            config_loader=self.config,
            embedding_factory=lambda config: KeywordEmbeddings(),
        )
        prepare_staged_snapshot(workspace, layout)
        add_staged_to_candidate(
            workspace,
            confirm_api_cost=True,
            confirm_batches=1,
            config_loader=self.config,
            embedding_factory=lambda config: KeywordEmbeddings(),
        )
        return workspace

    def test_plan_is_read_only_and_reports_two_exact_paid_phases(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            layout, upload_id = self.prepare_source_state(root)

            plan = plan_migration(
                layout,
                [upload_id],
                migration_id=MIGRATION_ID,
            )

            self.assertEqual(plan.legacy_record_count, 1)
            self.assertEqual(plan.base_chunk_count, 1)
            self.assertEqual(plan.incremental_new_chunk_count, 1)
            self.assertEqual(plan.final_chunk_count, 2)
            self.assertEqual(plan.base_embedding_batches, 1)
            self.assertEqual(plan.incremental_embedding_batches, 1)
            self.assertEqual(plan.total_embedding_batches, 2)
            self.assertFalse(layout.migrations_dir.exists())
            self.assertIsNone(load_index_manifest(layout.vector_store_dir))
            self.assertEqual(self.record_ids(layout.vector_store_dir), {"legacy-id"})

    def test_prepare_copies_only_base_without_moving_source_files(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            layout, upload_id = self.prepare_source_state(root)

            workspace = prepare_migration(
                layout,
                [upload_id],
                migration_id=MIGRATION_ID,
            )

            self.assertTrue((workspace / "base_documents" / "base.md").is_file())
            self.assertFalse((workspace / "candidate_documents").exists())
            self.assertTrue((layout.documents_dir / "base.md").is_file())
            self.assertTrue(
                (layout.pending_uploads_dir / upload_id / "word-notes.docx").is_file()
            )
            self.assertFalse((workspace / CANDIDATE_VECTOR_STORE_DIRNAME).exists())

    def test_base_phase_refuses_wrong_cost_confirmation_before_config_load(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            layout, upload_id = self.prepare_source_state(root)
            workspace = prepare_migration(
                layout,
                [upload_id],
                migration_id=MIGRATION_ID,
            )
            config_loader = Mock(side_effect=self.config)

            with self.assertRaisesRegex(LegacyMigrationError, "费用确认"):
                build_base_candidate(
                    workspace,
                    confirm_api_cost=True,
                    confirm_batches=2,
                    config_loader=config_loader,
                    embedding_factory=lambda config: KeywordEmbeddings(),
                )

            config_loader.assert_not_called()
            self.assertFalse((workspace / CANDIDATE_VECTOR_STORE_DIRNAME).exists())
            self.assertEqual(self.record_ids(layout.vector_store_dir), {"legacy-id"})

    def test_base_phase_recalculates_batches_before_loading_paid_config(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            layout, upload_id = self.prepare_source_state(root)
            workspace = prepare_migration(
                layout,
                [upload_id],
                migration_id=MIGRATION_ID,
            )
            plan_path = workspace / "plan.json"
            payload = json.loads(plan_path.read_text(encoding="utf-8"))
            payload["base_embedding_batches"] = 0
            payload["total_embedding_batches"] = 1
            plan_path.write_text(
                json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )
            config_loader = Mock(side_effect=self.config)

            with self.assertRaisesRegex(LegacyMigrationError, "切分结果"):
                build_base_candidate(
                    workspace,
                    confirm_api_cost=True,
                    confirm_batches=0,
                    config_loader=config_loader,
                    embedding_factory=lambda config: KeywordEmbeddings(),
                )

            config_loader.assert_not_called()
            self.assertFalse((workspace / CANDIDATE_VECTOR_STORE_DIRNAME).exists())

    def test_builds_and_validates_candidate_without_touching_legacy_index(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            layout, upload_id = self.prepare_source_state(root)

            workspace = self.build_complete_candidate(layout, upload_id)
            plan = validate_candidate(workspace)

            self.assertEqual(plan.final_chunk_count, 2)
            self.assertEqual(
                len(self.record_ids(workspace / CANDIDATE_VECTOR_STORE_DIRNAME)),
                2,
            )
            self.assertIsNotNone(
                load_index_manifest(workspace / CANDIDATE_VECTOR_STORE_DIRNAME)
            )
            self.assertEqual(self.record_ids(layout.vector_store_dir), {"legacy-id"})
            self.assertIsNone(load_index_manifest(layout.vector_store_dir))
            self.assertEqual(
                {path.name for path in layout.documents_dir.iterdir()},
                {"base.md"},
            )

    def test_incremental_failure_restores_verified_base_candidate(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            layout, upload_id = self.prepare_source_state(root)
            workspace = prepare_migration(
                layout,
                [upload_id],
                migration_id=MIGRATION_ID,
            )
            build_base_candidate(
                workspace,
                confirm_api_cost=True,
                confirm_batches=1,
                config_loader=self.config,
                embedding_factory=lambda config: KeywordEmbeddings(),
            )
            prepare_staged_snapshot(workspace, layout)
            base_ids = self.record_ids(workspace / CANDIDATE_VECTOR_STORE_DIRNAME)

            with self.assertRaisesRegex(RuntimeError, "simulated"):
                add_staged_to_candidate(
                    workspace,
                    confirm_api_cost=True,
                    confirm_batches=1,
                    config_loader=self.config,
                    embedding_factory=lambda config: FailingEmbeddings(),
                )

            self.assertEqual(
                self.record_ids(workspace / CANDIDATE_VECTOR_STORE_DIRNAME),
                base_ids,
            )
            state = (workspace / "state.json").read_text(encoding="utf-8")
            self.assertIn('"phase": "staged_snapshot_ready"', state)

    def test_incremental_state_write_failure_restores_base_candidate(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            layout, upload_id = self.prepare_source_state(root)
            workspace = prepare_migration(
                layout,
                [upload_id],
                migration_id=MIGRATION_ID,
            )
            build_base_candidate(
                workspace,
                confirm_api_cost=True,
                confirm_batches=1,
                config_loader=self.config,
                embedding_factory=lambda config: KeywordEmbeddings(),
            )
            prepare_staged_snapshot(workspace, layout)
            base_ids = self.record_ids(workspace / CANDIDATE_VECTOR_STORE_DIRNAME)

            with patch(
                "app.migrate_legacy_index._write_state",
                side_effect=OSError("simulated state write failure"),
            ):
                with self.assertRaisesRegex(OSError, "simulated state"):
                    add_staged_to_candidate(
                        workspace,
                        confirm_api_cost=True,
                        confirm_batches=1,
                        config_loader=self.config,
                        embedding_factory=lambda config: KeywordEmbeddings(),
                    )

            self.assertEqual(
                self.record_ids(workspace / CANDIDATE_VECTOR_STORE_DIRNAME),
                base_ids,
            )
            state = (workspace / "state.json").read_text(encoding="utf-8")
            self.assertIn('"phase": "staged_snapshot_ready"', state)

    def test_promote_and_rollback_preserve_both_versions(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            layout, upload_id = self.prepare_source_state(root)
            workspace = self.build_complete_candidate(layout, upload_id)

            backup = promote_candidate(
                workspace,
                layout,
                confirm_service_stopped=True,
                confirm_migration_id=MIGRATION_ID,
            )

            self.assertTrue((backup / "documents" / "base.md").is_file())
            self.assertIsNone(load_index_manifest(backup / "vector_store"))
            self.assertEqual(
                {path.name for path in layout.documents_dir.iterdir()},
                {"base.md", "word-notes.docx"},
            )
            self.assertIsNotNone(load_index_manifest(layout.vector_store_dir))
            self.assertEqual(len(self.record_ids(layout.vector_store_dir)), 2)

            rollback_promotion(
                workspace,
                layout,
                confirm_service_stopped=True,
                confirm_migration_id=MIGRATION_ID,
            )

            self.assertEqual(
                {path.name for path in layout.documents_dir.iterdir()},
                {"base.md"},
            )
            self.assertIsNone(load_index_manifest(layout.vector_store_dir))
            self.assertEqual(self.record_ids(layout.vector_store_dir), {"legacy-id"})
            self.assertTrue(
                (workspace / ROLLED_BACK_CANDIDATE_DIRNAME / "documents").is_dir()
            )

    def test_promotion_refuses_when_production_documents_drift(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            layout, upload_id = self.prepare_source_state(root)
            workspace = self.build_complete_candidate(layout, upload_id)
            (layout.documents_dir / "base.md").write_text("changed", encoding="utf-8")

            with self.assertRaisesRegex(LegacyMigrationError, "快照内容发生变化"):
                promote_candidate(
                    workspace,
                    layout,
                    confirm_service_stopped=True,
                    confirm_migration_id=MIGRATION_ID,
                )

            self.assertEqual(self.record_ids(layout.vector_store_dir), {"legacy-id"})
            self.assertTrue((workspace / CANDIDATE_VECTOR_STORE_DIRNAME).is_dir())

    def test_cli_rejects_workspace_outside_migration_root(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            layout = self.make_layout(root)
            outside = root / "outside" / MIGRATION_ID
            outside.mkdir(parents=True)

            with patch(
                "app.migrate_legacy_index.MigrationLayout",
                return_value=layout,
            ):
                with patch("builtins.print") as output:
                    exit_code = main(["validate", str(outside)])

            self.assertEqual(exit_code, 2)
            self.assertIn("安全范围", str(output.call_args.args[0]))


if __name__ == "__main__":
    unittest.main()

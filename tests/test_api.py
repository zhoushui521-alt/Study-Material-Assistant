import asyncio
import json
import tempfile
import unittest
from datetime import datetime
from io import BytesIO
from pathlib import Path
from uuid import UUID
from unittest.mock import AsyncMock, Mock, patch

from fastapi.testclient import TestClient
from docx import Document as WordDocument
from langchain_core.documents import Document

from app.api import (
    REQUEST_ID_HEADER,
    app,
    get_document_job_service,
    get_material_manager,
    get_rag_service_provider,
    get_web_material_service,
    lifespan,
)
from app.document_jobs import (
    DocumentJobConflictError,
    DocumentJobRecord,
    DocumentJobService,
)
from app.langchain_rag import LangChainRAGError, NO_EVIDENCE_ANSWER, RAGAnswer
from app.evidence import Citation
from app.material_ingestion import (
    BatchStageResult,
    IndexSyncSummary,
    MaterialDeleteResult,
    MaterialManager,
    MaterialRollbackError,
    StagedMaterialFailure,
    StagedMaterial,
)
from app.operation_guard import OperationGuard, OperationPolicy
from app.rag_service import RAGService, RAGServiceInitializationError
from app.request_history import RequestHistoryWriter
from app.url_safety import UnsafeURLError
from app.web_materials import (
    WebMaterialFetchError,
    WebMaterialPreview,
    WebMaterialService,
)
from tests.auth_helpers import TEST_USER, clear_user_services, install_authenticated_user


def pending_document_job(
    *,
    filenames: tuple[str, ...] = ("notes.md",),
) -> DocumentJobRecord:
    return DocumentJobRecord(
        job_id="11111111-1111-4111-8111-111111111111",
        user_id="22222222-2222-4222-8222-222222222222",
        upload_ids=tuple(chr(97 + index) * 32 for index in range(len(filenames))),
        filenames=filenames,
        status="pending",
        progress=0,
        error_message=None,
        result=None,
        created_at="2026-08-21T00:00:00.000Z",
        started_at=None,
        finished_at=None,
        updated_at="2026-08-21T00:00:00.000Z",
    )


class APITests(unittest.TestCase):
    def setUp(self) -> None:
        self.temporary_directory = tempfile.TemporaryDirectory()
        app.dependency_overrides.clear()
        install_authenticated_user()
        app.state.rag_service = None
        app.state.document_job_service = None
        app.state.operation_guard = OperationGuard()
        app.state.request_history_writer = RequestHistoryWriter(
            Path(self.temporary_directory.name) / "requests.jsonl"
        )

    def tearDown(self) -> None:
        app.dependency_overrides.clear()
        clear_user_services()
        app.state.rag_service = None
        app.state.document_job_service = None
        self.temporary_directory.cleanup()

    def test_health_does_not_initialize_rag(self) -> None:
        with patch("app.api.create_rag_service") as create_service:
            with TestClient(app) as client:
                response = client.get("/health")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json(), {"status": "ok"})
        create_service.assert_not_called()

    def test_web_material_service_receives_fake_ip_compatibility_switch(self) -> None:
        request = Mock()
        manager = Mock(spec=MaterialManager)
        request.app.state.material_manager = manager

        with (
            patch(
                "app.api.proxy_fake_ip_compatibility_enabled",
                return_value=True,
            ),
            patch("app.api.WebMaterialService") as service_class,
        ):
            service = get_web_material_service(request, manager)

        service_class.assert_called_once_with(
            manager,
            allow_proxy_fake_ip=True,
        )
        self.assertIs(service, service_class.return_value)

    def test_docs_are_available_without_initializing_rag(self) -> None:
        with patch("app.api.create_rag_service") as create_service:
            with TestClient(app) as client:
                response = client.get("/docs")

        self.assertEqual(response.status_code, 200)
        self.assertIn("swagger-ui", response.text)
        create_service.assert_not_called()

    def test_web_page_and_assets_do_not_initialize_rag(self) -> None:
        with self.assertLogs("uvicorn.error", level="INFO") as captured:
            with patch("app.api.create_rag_service") as create_service:
                with TestClient(app) as client:
                    page = client.get("/")
                    stylesheet = client.get("/static/styles.css")
                    script = client.get("/static/app.js")
                    hero_image = client.get("/static/zhixing-hero.png")
                    study_objects_image = client.get("/static/zhixing-study-objects.png")

        self.assertEqual(page.status_code, 200)
        self.assertIn("text/html", page.headers["content-type"])
        self.assertIn("知行 | Zhixing AI Learning Companion", page.text)
        self.assertIn('id="ask-form"', page.text)
        self.assertIn('id="citation-list"', page.text)
        self.assertIn('id="tutor-form"', page.text)
        self.assertIn('id="learning-history-list"', page.text)
        self.assertIn('id="web-preview-form"', page.text)
        self.assertIn(".docx", page.text)
        self.assertIn(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            page.text,
        )
        self.assertIn("multiple", page.text)
        self.assertIn('name="files"', page.text)
        self.assertIn("预计最多 60 个 Embedding 批次", page.text)
        self.assertIn("default-src 'self'", page.headers["content-security-policy"])
        self.assertEqual(page.headers["referrer-policy"], "no-referrer")

        self.assertEqual(stylesheet.status_code, 200)
        self.assertIn("text/css", stylesheet.headers["content-type"])
        self.assertIn("@media (max-width: 720px)", stylesheet.text)

        self.assertEqual(script.status_code, 200)
        self.assertIn("javascript", script.headers["content-type"])
        self.assertIn('fetch("/health"', script.text)
        self.assertIn('authenticatedFetch("/api/ask"', script.text)
        self.assertIn('authenticatedFetch("/api/tutor/chat"', script.text)
        self.assertIn('authenticatedFetch("/api/materials/stage-batch"', script.text)
        self.assertIn('fetch("/api/auth/me"', script.text)
        self.assertNotIn('"user_id": tutorIdentity', script.text)
        self.assertIn('"/api/materials/batch/index"', script.text)
        self.assertIn("单次上限 60", script.text)
        self.assertIn("confirm_api_cost: true", script.text)
        self.assertIn("IntersectionObserver", script.text)
        self.assertIn("textContent", script.text)
        self.assertNotIn("innerHTML", script.text)
        self.assertEqual(hero_image.status_code, 200)
        self.assertIn("image/png", hero_image.headers["content-type"])
        self.assertEqual(study_objects_image.status_code, 200)
        self.assertIn("image/png", study_objects_image.headers["content-type"])
        create_service.assert_not_called()
        request_paths = [
            json.loads(record.getMessage())["path"] for record in captured.records
        ]
        self.assertEqual(
            request_paths,
            ["/", "/static", "/static", "/static", "/static"],
        )

    def test_ask_returns_answer_and_source_labels(self) -> None:
        service = Mock(spec=RAGService)
        service.ask.return_value = RAGAnswer(
            answer="RAG 会先检索相关资料。[S1]",
            sources=(
                Document(
                    page_content="RAG 会先检索相关资料。",
                    metadata={"source": "rag.md", "chunk_index": 2},
                ),
            ),
            citations=(
                Citation(
                    citation_id="S1",
                    evidence_id="a" * 64,
                    material_id="b" * 64,
                    chunk_id="c" * 64,
                    source="rag.md",
                    filename="rag.md",
                    page=None,
                    chunk_index=2,
                    excerpt="RAG 会先检索相关资料。",
                    locator="rag.md#chunk=2",
                ),
            ),
        )
        app.dependency_overrides[get_rag_service_provider] = lambda: lambda: service

        with self.assertLogs("uvicorn.error", level="INFO") as captured:
            with TestClient(app) as client:
                response = client.post(
                    "/api/ask",
                    json={"question": " RAG 是什么？ "},
                )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.json(),
            {
                "answer": "RAG 会先检索相关资料。[S1]",
                "sources": ["[rag.md · 第 2 段]"],
                "citations": [
                    {
                        "citation_id": "S1",
                        "evidence_id": "a" * 64,
                        "material_id": "b" * 64,
                        "chunk_id": "c" * 64,
                        "source": "rag.md",
                        "filename": "rag.md",
                        "page": None,
                        "chunk_index": 2,
                        "excerpt": "RAG 会先检索相关资料。",
                        "locator": "rag.md#chunk=2",
                    }
                ],
            },
        )
        service.ask.assert_called_once_with("RAG 是什么？")
        request_log = json.loads(captured.records[-1].getMessage())
        self.assertEqual(request_log["event"], "http_request_completed")
        self.assertEqual(request_log["method"], "POST")
        self.assertEqual(request_log["path"], "/api/ask")
        self.assertEqual(request_log["status_code"], 200)
        self.assertIsNone(request_log["error_category"])
        self.assertGreaterEqual(request_log["elapsed_ms"], 0)
        self.assertEqual(response.headers[REQUEST_ID_HEADER], request_log["request_id"])
        UUID(request_log["request_id"])
        datetime.fromisoformat(request_log["timestamp"].replace("Z", "+00:00"))
        self.assertNotIn("RAG 是什么", captured.records[-1].getMessage())
        self.assertNotIn("RAG 会先检索", captured.records[-1].getMessage())

    def test_no_evidence_answer_has_no_sources(self) -> None:
        service = Mock(spec=RAGService)
        service.ask.return_value = RAGAnswer(
            answer=NO_EVIDENCE_ANSWER,
            sources=(),
        )
        app.dependency_overrides[get_rag_service_provider] = lambda: lambda: service

        with TestClient(app) as client:
            response = client.post(
                "/api/ask",
                json={"question": "MySQL 的事务隔离级别有哪些？"},
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["answer"], NO_EVIDENCE_ANSWER)
        self.assertEqual(response.json()["sources"], [])
        self.assertEqual(response.json()["citations"], [])

    def test_whitespace_question_is_rejected_before_rag_call(self) -> None:
        with self.assertLogs("uvicorn.error", level="INFO") as captured:
            with patch("app.api.create_rag_service") as create_service:
                with TestClient(app) as client:
                    response = client.post("/api/ask", json={"question": "   "})

        self.assertEqual(response.status_code, 422)
        create_service.assert_not_called()
        request_log = json.loads(captured.records[-1].getMessage())
        self.assertEqual(request_log["status_code"], 422)
        self.assertEqual(request_log["error_category"], "request_validation")

    def test_too_long_question_is_rejected_before_rag_call(self) -> None:
        with patch("app.api.create_rag_service") as create_service:
            with TestClient(app) as client:
                response = client.post(
                    "/api/ask",
                    json={"question": "问" * 2001},
                )

        self.assertEqual(response.status_code, 422)
        create_service.assert_not_called()

    def test_missing_and_unexpected_fields_are_rejected(self) -> None:
        with patch("app.api.create_rag_service") as create_service:
            with TestClient(app) as client:
                missing = client.post("/api/ask", json={})
                unexpected = client.post(
                    "/api/ask",
                    json={"question": "RAG？", "api_key": "should-not-be-accepted"},
                )

        self.assertEqual(missing.status_code, 422)
        self.assertEqual(unexpected.status_code, 422)
        create_service.assert_not_called()

    def test_rag_failure_returns_generic_502_without_sensitive_detail(self) -> None:
        service = Mock(spec=RAGService)
        service.ask.side_effect = LangChainRAGError(
            "调用 Chat 模型失败：api_key=secret-value"
        )
        app.dependency_overrides[get_rag_service_provider] = lambda: lambda: service

        with self.assertLogs("uvicorn.error", level="INFO") as captured:
            with TestClient(app) as client:
                response = client.post("/api/ask", json={"question": "RAG 是什么？"})

        self.assertEqual(response.status_code, 502)
        self.assertEqual(response.json(), {"detail": "RAG 问答处理失败。"})
        self.assertNotIn("secret-value", response.text)
        request_log_text = captured.records[-1].getMessage()
        request_log = json.loads(request_log_text)
        self.assertEqual(request_log["error_category"], "rag_processing")
        self.assertNotIn("secret-value", request_log_text)

    def test_initialization_failure_returns_generic_503(self) -> None:
        with patch(
            "app.api.create_rag_service",
            side_effect=RAGServiceInitializationError(
                "初始化失败：BAILIAN_API_KEY=secret-value"
            ),
        ):
            with self.assertLogs("uvicorn.error", level="INFO") as captured:
                with TestClient(app) as client:
                    response = client.post(
                        "/api/ask",
                        json={"question": "RAG 是什么？"},
                    )

        self.assertEqual(response.status_code, 503)
        self.assertEqual(response.json(), {"detail": "RAG 服务暂不可用。"})
        self.assertNotIn("secret-value", response.text)
        request_log_text = captured.records[-1].getMessage()
        request_log = json.loads(request_log_text)
        self.assertEqual(request_log["error_category"], "rag_unavailable")
        self.assertNotIn("secret-value", request_log_text)

    def test_unmatched_route_does_not_log_raw_path_or_query(self) -> None:
        with self.assertLogs("uvicorn.error", level="INFO") as captured:
            with TestClient(app) as client:
                response = client.get(
                    "/secret-path-value",
                    params={"api_key": "secret-query-value"},
                )

        self.assertEqual(response.status_code, 404)
        request_log_text = captured.records[-1].getMessage()
        request_log = json.loads(request_log_text)
        self.assertEqual(request_log["path"], "unmatched")
        self.assertEqual(request_log["error_category"], "client_error")
        self.assertNotIn("secret-path-value", request_log_text)
        self.assertNotIn("secret-query-value", request_log_text)

    def test_unknown_browser_page_returns_safe_html_without_request_values(self) -> None:
        with TestClient(app) as client:
            response = client.get(
                "/private-course-name",
                params={"token": "secret-token"},
                headers={"Accept": "text/html"},
            )

        self.assertEqual(response.status_code, 404)
        self.assertIn("text/html", response.headers["content-type"])
        self.assertIn("这里没有可打开的页面", response.text)
        self.assertIn("default-src 'self'", response.headers["content-security-policy"])
        self.assertNotIn("private-course-name", response.text)
        self.assertNotIn("secret-token", response.text)

    def test_unknown_api_route_returns_json_404(self) -> None:
        with TestClient(app) as client:
            response = client.get(
                "/api/not-real",
                headers={"Accept": "application/json"},
            )

        self.assertEqual(response.status_code, 404)
        self.assertEqual(response.json(), {"detail": "页面或接口不存在。"})

    def test_ask_rate_limit_returns_retry_after_without_extra_rag_call(self) -> None:
        service = Mock(spec=RAGService)
        service.ask.return_value = RAGAnswer(answer="回答", sources=())
        app.dependency_overrides[get_rag_service_provider] = lambda: lambda: service
        app.state.operation_guard = OperationGuard(
            policies={
                "ask": OperationPolicy(
                    window_seconds=60,
                    max_calls_per_window=1,
                    max_units_per_operation=1,
                    max_units_per_process=10,
                )
            }
        )

        with TestClient(app) as client:
            accepted = client.post("/api/ask", json={"question": "第一次"})
            rejected = client.post("/api/ask", json={"question": "第二次"})

        self.assertEqual(accepted.status_code, 200)
        self.assertEqual(rejected.status_code, 429)
        self.assertEqual(rejected.headers["retry-after"], "60")
        self.assertEqual(rejected.json(), {"detail": "操作过于频繁，请稍后重试。"})
        self.assertEqual(service.ask.call_count, 1)

    def test_unexpected_failure_returns_safe_500_with_request_id(self) -> None:
        service = Mock(spec=RAGService)
        service.ask.side_effect = RuntimeError("api_key=secret-value")
        app.dependency_overrides[get_rag_service_provider] = lambda: lambda: service

        with self.assertLogs("uvicorn.error", level="INFO") as captured:
            with TestClient(app) as client:
                response = client.post(
                    "/api/ask",
                    json={"question": "RAG 是什么？"},
                )

        self.assertEqual(response.status_code, 500)
        self.assertEqual(response.json(), {"detail": "服务器内部错误。"})
        request_log_text = captured.records[-1].getMessage()
        request_log = json.loads(request_log_text)
        self.assertEqual(request_log["error_category"], "unhandled_exception")
        self.assertEqual(response.headers[REQUEST_ID_HEADER], request_log["request_id"])
        self.assertNotIn("secret-value", response.text)
        self.assertNotIn("secret-value", request_log_text)

    def test_each_response_gets_a_unique_request_id(self) -> None:
        with TestClient(app) as client:
            first = client.get("/health")
            second = client.get("/health")

        self.assertNotEqual(
            first.headers[REQUEST_ID_HEADER],
            second.headers[REQUEST_ID_HEADER],
        )

    def test_lifespan_closes_initialized_service(self) -> None:
        service = Mock(spec=RAGService)
        app.state.rag_services[TEST_USER.user_id] = service

        with TestClient(app) as client:
            self.assertEqual(client.get("/health").status_code, 200)

        service.close.assert_called_once_with()
        self.assertNotIn(TEST_USER.user_id, app.state.rag_services)

    def test_lifespan_closes_service_after_context_error(self) -> None:
        service = Mock(spec=RAGService)
        app.state.rag_services[TEST_USER.user_id] = service

        async def exit_with_error() -> None:
            with self.assertRaisesRegex(RuntimeError, "application failed"):
                async with lifespan(app):
                    raise RuntimeError("application failed")

        asyncio.run(exit_with_error())

        service.close.assert_called_once_with()
        self.assertNotIn(TEST_USER.user_id, app.state.rag_services)

    def test_stale_cleanup_failure_does_not_block_startup_or_leak_detail(self) -> None:
        previous_manager = app.state.material_manager
        manager = Mock(spec=MaterialManager)
        manager.cleanup_stale_pending_uploads.side_effect = RuntimeError(
            "pending path contains secret-value"
        )
        app.state.material_manager = manager
        try:
            with self.assertLogs("uvicorn.error", level="WARNING") as captured:
                with TestClient(app) as client:
                    response = client.get("/health")
        finally:
            app.state.material_manager = previous_manager

        self.assertEqual(response.status_code, 200)
        log_text = "\n".join(record.getMessage() for record in captured.records)
        self.assertIn("stale_pending_upload_cleanup_failed", log_text)
        self.assertNotIn("secret-value", log_text)

    def test_lists_materials_without_initializing_rag(self) -> None:
        manager = Mock(spec=MaterialManager)
        manager.list_materials.return_value = ()
        app.dependency_overrides[get_material_manager] = lambda: manager

        with patch("app.api.create_rag_service") as create_service:
            with TestClient(app) as client:
                response = client.get("/api/materials")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json(), {"materials": []})
        create_service.assert_not_called()

    def test_stages_upload_without_indexing_or_initializing_rag(self) -> None:
        manager = Mock(spec=MaterialManager)
        manager.stage_upload.return_value = StagedMaterial(
            upload_id="a" * 32,
            filename="notes.md",
            operation="add",
            size_bytes=12,
            sha256="b" * 64,
            document_units=1,
            chunk_count=2,
            embedding_batch_count=1,
            staged_at="2026-08-09T10:00:00Z",
        )
        app.dependency_overrides[get_material_manager] = lambda: manager

        with patch("app.api.create_rag_service") as create_service:
            with TestClient(app) as client:
                response = client.post(
                    "/api/materials/stage",
                    files={"file": ("notes.md", b"RAG notes", "text/markdown")},
                    data={"operation": "add"},
                )

        self.assertEqual(response.status_code, 202)
        self.assertEqual(response.json()["status"], "staged")
        self.assertEqual(response.json()["chunk_count"], 2)
        manager.stage_upload.assert_called_once()
        call = manager.stage_upload.call_args.kwargs
        self.assertEqual(call["filename"], "notes.md")
        self.assertEqual(call["operation"], "add")
        create_service.assert_not_called()

    def test_stages_real_docx_parser_through_api_without_indexing(self) -> None:
        root = Path(self.temporary_directory.name) / "docx-api"
        sync_index = Mock(return_value=IndexSyncSummary(1, 0, 0))
        manager = MaterialManager(
            documents_dir=root / "documents",
            pending_uploads_dir=root / "pending_uploads",
            pending_deletions_dir=root / "pending_deletions",
            sync_index=sync_index,
            delete_index=lambda filename: 0,
            estimate_index_batches=lambda chunks: 1,
        )
        app.dependency_overrides[get_material_manager] = lambda: manager
        stream = BytesIO()
        document = WordDocument()
        document.add_heading("RAG", level=1)
        document.add_paragraph("DOCX 通过 API 进入现有资料管线。")
        document.save(stream)

        with patch("app.api.create_rag_service") as create_service:
            with TestClient(app) as client:
                response = client.post(
                    "/api/materials/stage",
                    files={
                        "file": (
                            "notes.docx",
                            stream.getvalue(),
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document",
                        )
                    },
                    data={"operation": "add"},
                )

        self.assertEqual(response.status_code, 202)
        self.assertEqual(response.json()["document_units"], 2)
        self.assertEqual(response.json()["chunk_count"], 2)
        sync_index.assert_not_called()
        create_service.assert_not_called()

    def test_batch_stage_returns_success_and_failure_details(self) -> None:
        manager = Mock(spec=MaterialManager)
        manager.stage_upload_batch.return_value = BatchStageResult(
            staged=(
                StagedMaterial(
                    upload_id="a" * 32,
                    filename="notes.md",
                    operation="add",
                    size_bytes=12,
                    sha256="b" * 64,
                    document_units=1,
                    chunk_count=2,
                    embedding_batch_count=1,
                    staged_at="2026-08-19T10:00:00Z",
                ),
            ),
            failures=(
                StagedMaterialFailure(
                    filename="broken.docx",
                    reason="DOCX 文件已损坏。",
                ),
            ),
        )
        app.dependency_overrides[get_material_manager] = lambda: manager

        with TestClient(app) as client:
            response = client.post(
                "/api/materials/stage-batch",
                files=[
                    ("files", ("notes.md", b"notes", "text/markdown")),
                    ("files", ("broken.docx", b"broken", "application/zip")),
                ],
                data={"operation": "add"},
            )

        self.assertEqual(response.status_code, 207)
        self.assertEqual(response.json()["status"], "partial")
        self.assertEqual(response.json()["staged_count"], 1)
        self.assertEqual(response.json()["failed_count"], 1)
        self.assertEqual(
            response.json()["failures"],
            [{"filename": "broken.docx", "reason": "DOCX 文件已损坏。"}],
        )
        uploads = manager.stage_upload_batch.call_args.args[0]
        self.assertEqual([item.filename for item in uploads], ["notes.md", "broken.docx"])

    def test_batch_stage_all_failures_returns_structured_422(self) -> None:
        manager = Mock(spec=MaterialManager)
        manager.stage_upload_batch.return_value = BatchStageResult(
            staged=(),
            failures=(
                StagedMaterialFailure("a.docx", "文件损坏。"),
                StagedMaterialFailure("b.doc", "旧版 .doc 暂不支持。"),
            ),
        )
        app.dependency_overrides[get_material_manager] = lambda: manager

        with TestClient(app) as client:
            response = client.post(
                "/api/materials/stage-batch",
                files=[
                    ("files", ("a.docx", b"a", "application/zip")),
                    ("files", ("b.doc", b"b", "application/msword")),
                ],
                data={"operation": "add"},
            )

        self.assertEqual(response.status_code, 422)
        self.assertEqual(response.json()["status"], "failed")
        self.assertEqual(response.json()["failed_count"], 2)

    def test_previews_web_material_without_indexing_or_initializing_rag(self) -> None:
        service = Mock(spec=WebMaterialService)
        service.preview = AsyncMock(
            return_value=WebMaterialPreview(
                upload_id="c" * 32,
                filename="web-example-com-123456789abc.md",
                operation="add",
                requested_url="https://example.com/start",
                canonical_url="https://example.com/rag",
                title="RAG Guide",
                crawled_at="2026-08-10T09:30:00Z",
                content_sha256="d" * 64,
                markdown="# RAG\n\n先检索，再生成。",
                redirect_count=1,
                size_bytes=320,
                document_units=1,
                chunk_count=2,
                embedding_batch_count=1,
            )
        )
        app.dependency_overrides[get_web_material_service] = lambda: service

        with patch("app.api.create_rag_service") as create_service:
            with TestClient(app) as client:
                response = client.post(
                    "/api/web-materials/preview",
                    json={"url": "https://example.com/start", "operation": "add"},
                )

        self.assertEqual(response.status_code, 202)
        self.assertEqual(response.json()["status"], "staged")
        self.assertEqual(response.json()["canonical_url"], "https://example.com/rag")
        service.preview.assert_awaited_once_with(
            "https://example.com/start",
            operation="add",
        )
        create_service.assert_not_called()

    def test_web_preview_rejects_extra_fields_before_service_call(self) -> None:
        service = Mock(spec=WebMaterialService)
        service.preview = AsyncMock()
        app.dependency_overrides[get_web_material_service] = lambda: service

        with TestClient(app) as client:
            response = client.post(
                "/api/web-materials/preview",
                json={
                    "url": "https://example.com/",
                    "operation": "add",
                    "cookie": "secret",
                },
            )

        self.assertEqual(response.status_code, 422)
        service.preview.assert_not_awaited()

    def test_web_preview_maps_unsafe_url_to_422_without_logging_url(self) -> None:
        service = Mock(spec=WebMaterialService)
        service.preview = AsyncMock(
            side_effect=UnsafeURLError("URL 只能解析到公网 IP 地址。")
        )
        app.dependency_overrides[get_web_material_service] = lambda: service
        private_url = "http://127.0.0.1/private-token"

        with TestClient(app) as client:
            response = client.post(
                "/api/web-materials/preview",
                json={"url": private_url, "operation": "add"},
            )

        self.assertEqual(response.status_code, 422)
        log_text = (
            Path(self.temporary_directory.name) / "requests.jsonl"
        ).read_text(encoding="utf-8")
        self.assertNotIn(private_url, log_text)
        self.assertNotIn("private-token", log_text)
        record = json.loads(log_text.splitlines()[-1])
        self.assertEqual(record["path"], "/api/web-materials/preview")
        self.assertEqual(record["error_category"], "web_preview")

    def test_web_preview_returns_generic_fetch_error_without_raw_detail(self) -> None:
        service = Mock(spec=WebMaterialService)
        service.preview = AsyncMock(
            side_effect=WebMaterialFetchError("upstream contained secret-value")
        )
        app.dependency_overrides[get_web_material_service] = lambda: service

        with TestClient(app) as client:
            response = client.post(
                "/api/web-materials/preview",
                json={"url": "https://example.com/", "operation": "add"},
            )

        self.assertEqual(response.status_code, 502)
        self.assertEqual(response.json(), {"detail": "公开网页抓取失败。"})
        self.assertNotIn("secret-value", response.text)

    def test_web_preview_rate_limit_returns_retry_after(self) -> None:
        service = Mock(spec=WebMaterialService)
        service.preview = AsyncMock(
            return_value=WebMaterialPreview(
                upload_id="c" * 32,
                filename="web-example-com-123456789abc.md",
                operation="add",
                requested_url="https://example.com/",
                canonical_url="https://example.com/",
                title="Example",
                crawled_at="2026-08-10T09:30:00Z",
                content_sha256="d" * 64,
                markdown="Example",
                redirect_count=0,
                size_bytes=200,
                document_units=1,
                chunk_count=1,
                embedding_batch_count=1,
            )
        )
        app.dependency_overrides[get_web_material_service] = lambda: service
        app.state.operation_guard = OperationGuard(
            policies={
                "web_preview": OperationPolicy(
                    window_seconds=60,
                    max_calls_per_window=1,
                )
            }
        )

        with TestClient(app) as client:
            accepted = client.post(
                "/api/web-materials/preview",
                json={"url": "https://example.com/", "operation": "add"},
            )
            rejected = client.post(
                "/api/web-materials/preview",
                json={"url": "https://example.com/", "operation": "add"},
            )

        self.assertEqual(accepted.status_code, 202)
        self.assertEqual(rejected.status_code, 429)
        self.assertEqual(rejected.headers["retry-after"], "60")
        self.assertEqual(service.preview.await_count, 1)

    def test_index_requires_literal_cost_confirmation(self) -> None:
        service = Mock(spec=DocumentJobService)
        service.enqueue = AsyncMock()
        app.dependency_overrides[get_document_job_service] = lambda: service

        with TestClient(app) as client:
            response = client.post(
                f"/api/materials/{'a' * 32}/index",
                json={"confirm_api_cost": False},
            )

        self.assertEqual(response.status_code, 422)
        service.enqueue.assert_not_called()

    def test_confirmed_index_returns_pending_job_without_running_pipeline(self) -> None:
        service = Mock(spec=DocumentJobService)
        service.enqueue = AsyncMock(return_value=pending_document_job())
        app.dependency_overrides[get_document_job_service] = lambda: service

        with TestClient(app) as client:
            response = client.post(
                f"/api/materials/{'a' * 32}/index",
                json={"confirm_api_cost": True},
            )

        self.assertEqual(response.status_code, 202)
        self.assertEqual(response.json()["status"], "pending")
        self.assertEqual(
            response.json()["job_id"],
            "11111111-1111-4111-8111-111111111111",
        )
        service.enqueue.assert_awaited_once_with(
            TEST_USER.user_id, ("a" * 32,)
        )

    def test_confirmed_batch_index_returns_one_pending_job(self) -> None:
        upload_ids = ["a" * 32, "b" * 32]
        service = Mock(spec=DocumentJobService)
        service.enqueue = AsyncMock(
            return_value=pending_document_job(
                filenames=("chapter.pdf", "notes.docx")
            )
        )
        app.dependency_overrides[get_document_job_service] = lambda: service

        with TestClient(app) as client:
            response = client.post(
                "/api/materials/batch/index",
                json={"upload_ids": upload_ids, "confirm_api_cost": True},
            )

        self.assertEqual(response.status_code, 202)
        self.assertEqual(response.json()["filenames"], ["chapter.pdf", "notes.docx"])
        service.enqueue.assert_awaited_once_with(TEST_USER.user_id, upload_ids)

    def test_duplicate_active_index_job_returns_conflict(self) -> None:
        upload_ids = ["a" * 32, "b" * 32]
        service = Mock(spec=DocumentJobService)
        service.enqueue = AsyncMock(
            side_effect=DocumentJobConflictError(
                "这些暂存资料已经有等待或正在处理的任务。"
            )
        )
        app.dependency_overrides[get_document_job_service] = lambda: service

        with TestClient(app) as client:
            response = client.post(
                "/api/materials/batch/index",
                json={"upload_ids": upload_ids, "confirm_api_cost": True},
            )

        self.assertEqual(response.status_code, 409)
        self.assertIn("已经有等待", response.json()["detail"])

    def test_queries_completed_document_job(self) -> None:
        record = pending_document_job()
        completed = DocumentJobRecord(
            **{
                **record.__dict__,
                "status": "completed",
                "progress": 100,
                "result": {
                    "filenames": ["notes.md"],
                    "added": 2,
                    "deleted": 0,
                    "unchanged": 141,
                    "cleanup_pending": False,
                },
                "started_at": "2026-08-21T00:00:01.000Z",
                "finished_at": "2026-08-21T00:00:02.000Z",
            }
        )
        service = Mock(spec=DocumentJobService)
        service.get_job = AsyncMock(return_value=completed)
        app.dependency_overrides[get_document_job_service] = lambda: service

        with TestClient(app) as client:
            response = client.get(
                "/api/jobs/11111111-1111-4111-8111-111111111111"
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json()["status"], "completed")
        self.assertEqual(response.json()["result"]["added"], 2)

    def test_delete_requires_confirmation_and_invalidates_rag(self) -> None:
        manager = Mock(spec=MaterialManager)
        manager.delete_material.return_value = MaterialDeleteResult(
            filename="notes.md",
            deleted_records=2,
            cleanup_pending=False,
        )
        app.dependency_overrides[get_material_manager] = lambda: manager

        with TestClient(app) as client:
            rejected = client.request(
                "DELETE",
                "/api/materials/notes.md",
                json={"confirm_delete": False},
            )
            accepted = client.request(
                "DELETE",
                "/api/materials/notes.md",
                json={"confirm_delete": True},
            )

        self.assertEqual(rejected.status_code, 422)
        self.assertEqual(accepted.status_code, 200)
        self.assertEqual(accepted.json()["deleted_records"], 2)
        manager.delete_material.assert_called_once_with("notes.md")
        log_text = (
            Path(self.temporary_directory.name) / "requests.jsonl"
        ).read_text(encoding="utf-8")
        self.assertNotIn("notes.md", log_text)
        self.assertEqual(
            json.loads(log_text.splitlines()[-1])["path"],
            "/api/materials/{filename}",
        )

    def test_delete_uncertain_rollback_returns_503_and_invalidates_rag(self) -> None:
        manager = Mock(spec=MaterialManager)
        manager.delete_material.side_effect = MaterialRollbackError(
            "sensitive internal detail"
        )
        app.dependency_overrides[get_material_manager] = lambda: manager
        cached_service = Mock(spec=RAGService)
        app.state.rag_services[TEST_USER.user_id] = cached_service

        with TestClient(app) as client:
            response = client.request(
                "DELETE",
                "/api/materials/notes.md",
                json={"confirm_delete": True},
            )

        self.assertEqual(response.status_code, 503)
        self.assertNotIn("sensitive internal detail", response.text)
        cached_service.close.assert_called_once_with()
        self.assertNotIn(TEST_USER.user_id, app.state.rag_services)

    def test_request_metadata_is_persisted_without_question_or_answer(self) -> None:
        service = Mock(spec=RAGService)
        service.ask.return_value = RAGAnswer(answer="安全回答", sources=())
        app.dependency_overrides[get_rag_service_provider] = lambda: lambda: service

        with TestClient(app) as client:
            response = client.post("/api/ask", json={"question": "私密问题"})

        self.assertEqual(response.status_code, 200)
        log_text = (
            Path(self.temporary_directory.name) / "requests.jsonl"
        ).read_text(encoding="utf-8")
        self.assertNotIn("私密问题", log_text)
        self.assertNotIn("安全回答", log_text)
        record = json.loads(log_text.splitlines()[-1])
        self.assertEqual(record["path"], "/api/ask")


if __name__ == "__main__":
    unittest.main()

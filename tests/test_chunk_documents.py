import hashlib
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import Mock, patch

from docx import Document as WordDocument
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.chunk_documents import (
    DocumentLoadError,
    WebSourceMetadata,
    build_chunks,
    encode_web_source_marker,
    extract_pdf_pages,
    load_documents,
    load_material_units,
    split_text,
)
from app.evidence import evidence_from_document
from app.langchain_store import chunks_to_documents


def write_docx(
    path: Path,
    *,
    paragraphs: tuple[str, ...] = (),
    headings: tuple[tuple[str, int], ...] = (),
    table_rows: tuple[tuple[str, ...], ...] = (),
) -> None:
    document = WordDocument()
    for text, level in headings:
        document.add_heading(text, level=level)
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for column_index, value in enumerate(row):
                table.cell(row_index, column_index).text = value
    document.save(path)


def write_text_pdf(path: Path, text: str) -> None:
    """创建一个带 Helvetica 文字层的最小测试 PDF。"""
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=300)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
    )
    content = DecodedStreamObject()
    escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content.set_data(f"BT /F1 12 Tf 72 200 Td ({escaped_text}) Tj ET".encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(content)
    with path.open("wb") as stream:
        writer.write(stream)


class ChunkDocumentsTests(unittest.TestCase):
    def test_docx_paragraph_enters_material_units_and_existing_chunker(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            write_docx(
                documents_dir / "rag.docx",
                paragraphs=("Embedding 是将文本转换为向量表示的方法。",),
            )

            units = load_material_units(documents_dir)
            chunks = build_chunks(units)

        self.assertEqual(len(units), 1)
        self.assertEqual(units[0].material.source_type, "docx")
        self.assertEqual(units[0].paragraph_index, 1)
        self.assertEqual(units[0].source, "rag.docx · 第 1 段")
        self.assertIn("Embedding", chunks[0].content)

    def test_docx_heading_is_preserved_as_section_for_following_paragraph(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            path = documents_dir / "rag.docx"
            document = WordDocument()
            document.add_heading("检索", level=1)
            document.add_heading("向量召回", level=2)
            document.add_heading("候选资料", level=3)
            document.add_paragraph("Retriever 负责从知识库寻找候选资料。")
            document.save(path)

            units = load_material_units(documents_dir)

        self.assertEqual(units[-1].section, "检索 > 向量召回 > 候选资料")
        self.assertEqual(units[-1].paragraph_index, 4)

    def test_docx_table_becomes_searchable_text_in_document_order(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            write_docx(
                documents_dir / "concepts.docx",
                headings=(("RAG 基础", 1),),
                table_rows=(("概念", "定义"), ("RAG", "检索增强生成")),
            )

            units = load_material_units(documents_dir)
            chunks = build_chunks(units)

        table_unit = next(unit for unit in units if unit.table_index is not None)
        table_chunk = next(chunk for chunk in chunks if chunk.table_index is not None)
        self.assertEqual(table_unit.section, "RAG 基础")
        self.assertEqual(table_unit.source, "concepts.docx · 第 1 个表格")
        self.assertIn("概念 | 定义", table_chunk.content)
        self.assertIn("RAG | 检索增强生成", table_chunk.content)

    def test_docx_metadata_reaches_evidence_without_fake_page(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            write_docx(
                documents_dir / "rag.docx",
                headings=(("检索基础", 1),),
                paragraphs=("向量检索会返回候选证据。",),
            )
            chunk = build_chunks(load_material_units(documents_dir))[-1]

        document = chunks_to_documents([chunk])[0]
        evidence = evidence_from_document(document, "S1")
        self.assertIsNone(chunk.page)
        self.assertNotIn("page", document.metadata)
        self.assertEqual(document.metadata["paragraph_index"], 2)
        self.assertEqual(document.metadata["section"], "检索基础")
        self.assertIsNone(evidence.page)
        self.assertIn("section=%E6%A3%80%E7%B4%A2%E5%9F%BA%E7%A1%80", evidence.locator)
        self.assertIn("paragraph=2&chunk=1", evidence.locator)

    def test_empty_docx_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            write_docx(documents_dir / "empty.docx")

            with self.assertRaisesRegex(DocumentLoadError, "没有可建立索引"):
                load_material_units(documents_dir)

    def test_corrupted_docx_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            (documents_dir / "broken.docx").write_bytes(b"PK\x03\x04broken")

            with self.assertRaisesRegex(DocumentLoadError, "不是有效|损坏"):
                load_material_units(documents_dir)

    def test_zip_without_required_ooxml_parts_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            path = documents_dir / "fake.docx"
            with zipfile.ZipFile(path, "w") as archive:
                archive.writestr("readme.txt", "not a Word package")

            with self.assertRaisesRegex(DocumentLoadError, "缺少必要的 OOXML 结构"):
                load_material_units(documents_dir)

    def test_docx_repeated_parse_keeps_stable_identity(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            write_docx(documents_dir / "rag.docx", paragraphs=("稳定内容",))

            first = build_chunks(load_material_units(documents_dir))[0]
            second = build_chunks(load_material_units(documents_dir))[0]

        self.assertEqual(first.material_id, second.material_id)
        self.assertEqual(first.material_content_hash, second.material_content_hash)
        self.assertEqual(first.chunk_id, second.chunk_id)

    def test_identical_docx_paragraphs_have_distinct_chunk_identity(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            write_docx(
                documents_dir / "rag.docx",
                paragraphs=("相同内容", "相同内容"),
            )

            chunks = build_chunks(load_material_units(documents_dir))

        self.assertEqual([chunk.paragraph_index for chunk in chunks], [1, 2])
        self.assertNotEqual(chunks[0].chunk_id, chunks[1].chunk_id)

    def test_docx_content_change_keeps_material_identity_and_changes_hashes(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            path = documents_dir / "rag.docx"
            write_docx(path, paragraphs=("原始内容",))
            original = build_chunks(load_material_units(documents_dir))[0]

            write_docx(path, paragraphs=("修改后的内容",))
            changed = build_chunks(load_material_units(documents_dir))[0]

        self.assertEqual(original.material_id, changed.material_id)
        self.assertNotEqual(original.material_content_hash, changed.material_content_hash)
        self.assertNotEqual(original.content_hash, changed.content_hash)
        self.assertNotEqual(original.chunk_id, changed.chunk_id)

    def test_load_documents_uses_web_url_as_source_and_removes_marker(self) -> None:
        body = "# RAG\n\n正文"
        metadata = WebSourceMetadata(
            canonical_url="https://example.com/rag",
            title="RAG Guide",
            crawled_at="2026-08-10T09:30:00Z",
            content_sha256=hashlib.sha256(body.encode("utf-8")).hexdigest(),
        )
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            (documents_dir / "web-example.md").write_text(
                f"{encode_web_source_marker(metadata)}\n{body}",
                encoding="utf-8",
            )

            documents = load_documents(documents_dir)
            units = load_material_units(documents_dir)

        self.assertEqual(
            documents,
            [("web-example.md · 网页：https://example.com/rag", "# RAG\n\n正文")],
        )
        self.assertEqual(units[0].material.source_type, "web")
        self.assertEqual(units[0].material.canonical_url, "https://example.com/rag")
        self.assertIsNone(units[0].page)

    def test_rejects_tampered_web_material_body(self) -> None:
        original_body = "# RAG\n\n原始正文"
        metadata = WebSourceMetadata(
            canonical_url="https://example.com/rag",
            title="RAG Guide",
            crawled_at="2026-08-10T09:30:00Z",
            content_sha256=hashlib.sha256(original_body.encode("utf-8")).hexdigest(),
        )
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            (documents_dir / "web-example.md").write_text(
                f"{encode_web_source_marker(metadata)}\n# RAG\n\n篡改正文",
                encoding="utf-8",
            )

            with self.assertRaisesRegex(DocumentLoadError, "正文与来源元数据不一致"):
                load_documents(documents_dir)

    def test_rejects_web_marker_with_invalid_port(self) -> None:
        metadata = WebSourceMetadata(
            canonical_url="https://example.com:bad/rag",
            title="RAG Guide",
            crawled_at="2026-08-10T09:30:00Z",
            content_sha256="a" * 64,
        )

        with self.assertRaisesRegex(DocumentLoadError, "元数据无效"):
            encode_web_source_marker(metadata)

    def test_rejects_web_marker_with_non_default_port(self) -> None:
        metadata = WebSourceMetadata(
            canonical_url="https://example.com:80/rag",
            title="RAG Guide",
            crawled_at="2026-08-10T09:30:00Z",
            content_sha256="a" * 64,
        )

        with self.assertRaisesRegex(DocumentLoadError, "元数据无效"):
            encode_web_source_marker(metadata)

    def test_split_text_splits_long_text(self) -> None:
        chunks = split_text("a" * 181, chunk_size=180)
        self.assertEqual(chunks, ["a" * 180, "a"])

    def test_build_chunks_keeps_source_information(self) -> None:
        chunks = build_chunks([("notes.md", "RAG 会先检索资料。")])
        self.assertEqual(chunks[0].source, "notes.md")
        self.assertEqual(chunks[0].index, 1)

    def test_repeated_parse_keeps_material_and_chunk_identity_stable(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            (documents_dir / "notes.md").write_text("RAG 会先检索资料。", encoding="utf-8")

            first = build_chunks(load_material_units(documents_dir))[0]
            second = build_chunks(load_material_units(documents_dir))[0]

        self.assertEqual(first.material_id, second.material_id)
        self.assertEqual(first.chunk_id, second.chunk_id)
        self.assertEqual(first.content_hash, second.content_hash)

    def test_content_change_keeps_material_identity_but_changes_content_identity(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            path = documents_dir / "notes.md"
            path.write_text("RAG 会先检索资料。", encoding="utf-8")
            original = build_chunks(load_material_units(documents_dir))[0]

            path.write_text("RAG 会先检索资料并生成回答。", encoding="utf-8")
            changed = build_chunks(load_material_units(documents_dir))[0]

        self.assertEqual(original.material_id, changed.material_id)
        self.assertNotEqual(original.material_content_hash, changed.material_content_hash)
        self.assertNotEqual(original.content_hash, changed.content_hash)
        self.assertNotEqual(original.chunk_id, changed.chunk_id)

    def test_load_documents_extracts_pdf_text_and_page_number(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            write_text_pdf(documents_dir / "rag.pdf", "RAG PDF content")

            documents = dict(load_documents(documents_dir))

            source = "rag.pdf · 第 1 页"
            self.assertIn(source, documents)
        self.assertIn("RAG PDF content", documents[source])

    def test_pdf_page_is_structured_and_text_material_does_not_fake_page(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            write_text_pdf(documents_dir / "rag.pdf", "RAG PDF content")
            (documents_dir / "notes.txt").write_text("RAG text", encoding="utf-8")

            chunks = build_chunks(load_material_units(documents_dir))

        pdf_chunk = next(chunk for chunk in chunks if chunk.filename == "rag.pdf")
        text_chunk = next(chunk for chunk in chunks if chunk.filename == "notes.txt")
        self.assertEqual(pdf_chunk.page, 1)
        self.assertEqual(pdf_chunk.source_type, "pdf")
        self.assertIsNone(text_chunk.page)
        self.assertEqual(text_chunk.source_type, "text")

    @patch("app.chunk_documents.pdfplumber.open")
    def test_structured_pdf_page_keeps_original_number_after_blank_page(
        self,
        mock_open: Mock,
    ) -> None:
        blank_page = Mock()
        blank_page.extract_text.return_value = ""
        content_page = Mock()
        content_page.extract_text.return_value = "second page"
        mock_open.return_value.__enter__.return_value.pages = [
            blank_page,
            content_page,
        ]
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            (documents_dir / "rag.pdf").write_bytes(b"test-pdf-content")

            units = load_material_units(documents_dir)

        self.assertEqual(len(units), 1)
        self.assertEqual(units[0].page, 2)
        self.assertEqual(units[0].source, "rag.pdf · 第 2 页")

    @patch("app.chunk_documents.pdfplumber.open")
    def test_extract_pdf_pages_preserves_chinese_text(self, mock_open: Mock) -> None:
        page = Mock()
        page.extract_text.return_value = "第 8 章 多元分布"
        mock_open.return_value.__enter__.return_value.pages = [page]

        documents = extract_pdf_pages(Path("documents.pdf"))

        self.assertEqual(
            documents,
            [("documents.pdf · 第 1 页", "第 8 章 多元分布")],
        )

    @patch("app.chunk_documents.pdfplumber.open")
    def test_extract_pdf_pages_rejects_page_count_over_limit(
        self,
        mock_open: Mock,
    ) -> None:
        pages = [Mock(), Mock()]
        mock_open.return_value.__enter__.return_value.pages = pages

        with self.assertRaisesRegex(DocumentLoadError, "超过 1 页"):
            extract_pdf_pages(Path("documents.pdf"), max_pages=1)

        for page in pages:
            page.extract_text.assert_not_called()

    @patch("app.chunk_documents.pdfplumber.open")
    def test_extract_pdf_pages_rejects_extracted_text_over_limit(
        self,
        mock_open: Mock,
    ) -> None:
        page = Mock()
        page.extract_text.return_value = "123456"
        mock_open.return_value.__enter__.return_value.pages = [page]

        with self.assertRaisesRegex(DocumentLoadError, "提取文字超过安全限制"):
            extract_pdf_pages(
                Path("documents.pdf"),
                max_total_characters=5,
            )

    def test_load_documents_rejects_total_text_over_limit(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            documents_dir = Path(directory)
            (documents_dir / "notes.md").write_text("123456", encoding="utf-8")

            with self.assertRaisesRegex(DocumentLoadError, "提取文字超过安全限制"):
                load_documents(documents_dir, max_total_characters=5)

    def test_load_documents_rejects_pdf_without_text_layer(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            pdf_path = Path(directory) / "scan.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=300, height=300)
            with pdf_path.open("wb") as stream:
                writer.write(stream)

            with self.assertRaisesRegex(DocumentLoadError, "需要先进行 OCR"):
                load_documents(Path(directory))

    def test_load_documents_reports_encrypted_pdf(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            pdf_path = Path(directory) / "secret.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=300, height=300)
            writer.encrypt("secret")
            with pdf_path.open("wb") as stream:
                writer.write(stream)

            with self.assertRaisesRegex(DocumentLoadError, "暂不支持加密文件"):
                load_documents(Path(directory))

    def test_load_documents_reports_corrupted_pdf(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            pdf_path = Path(directory) / "broken.pdf"
            pdf_path.write_bytes(b"not a valid PDF")

            with self.assertRaisesRegex(DocumentLoadError, "文件可能已损坏"):
                load_documents(Path(directory))


if __name__ == "__main__":
    unittest.main()
